import re
from datetime import date

from docx import Document

from app.db.repositories import ResumeRepository, TailoredResumeRepository
from app.models.resume import Resume
from app.models.tailored_resume import TailoredResume
from app.services.llm_service import JD_PROMPT_LIMIT, RESUME_PROMPT_CHAR_BUDGET, LLMService
from app.utils.file_utils import RESUMES_TAILORED_DIR, ensure_app_directories, unique_path

JD_MIN_LENGTH = 100
HEADER_SCAN_LINES = 6

JOB_TITLE_MAX_LENGTH = 25
COMPANY_MAX_LENGTH = 20
JOB_TITLE_REMOVE_WORDS = {"senior", "lead", "mid", "jr", "remote", "contract", "hybrid"}
COMPANY_REMOVE_WORDS = {"inc", "llc", "corp", "ltd", "co", "technologies", "solutions"}

CONTACT_LINE_PATTERN = re.compile(
    r"@|linkedin\.com|github\.com|\(\d{3}\)|\d{3}[-.\s]\d{3}[-.\s]\d{4}", re.IGNORECASE
)


class ShortJobDescriptionError(Exception):
    pass


class EmptyResumeError(Exception):
    pass


def tailor_resume(
    job_id: int,
    company: str,
    resume_path: str,
    jd_text: str,
    job_title: str,
    role_description: str,
    profile,
    score: int | None = None,
) -> TailoredResume:
    if not jd_text or len(jd_text.strip()) < JD_MIN_LENGTH:
        raise ShortJobDescriptionError(
            f"Job description is too short to tailor a resume against "
            f"(needs at least {JD_MIN_LENGTH} characters)."
        )

    document = Document(resume_path)
    paragraphs = document.paragraphs
    if not paragraphs:
        raise EmptyResumeError("This resume has no readable paragraphs to tailor.")

    header_end_index = _detect_header_end(paragraphs)
    original_lines = [paragraph.text for paragraph in paragraphs]
    numbered_text, included_count = _numbered_text_within_budget(
        original_lines, RESUME_PROMPT_CHAR_BUDGET
    )

    tailored_by_line_number = LLMService().tailor(
        numbered_text,
        jd_text[:JD_PROMPT_LIMIT],
        job_title,
        role_description,
        profile,
        header_line_count=header_end_index + 1,
    )

    final_lines = _apply_tailored_lines(paragraphs, header_end_index, tailored_by_line_number)

    ensure_app_directories()
    file_name = _build_file_name(company, job_title)
    destination = unique_path(RESUMES_TAILORED_DIR, file_name)
    document.save(str(destination))

    resume_id = ResumeRepository().add(
        Resume(
            file_name=destination.name,
            file_path=str(destination),
            is_default=False,
            notes=f"Tailored for {company} — {job_title}",
        )
    )

    tailored = TailoredResume(
        job_id=job_id,
        resume_id=resume_id,
        company=company,
        job_title=job_title,
        score=score,
        file_name=destination.name,
        file_path=str(destination),
        tailored_text="\n".join(final_lines),
        source_resume_path=str(resume_path),
    )
    tailored.id = TailoredResumeRepository().save(tailored)
    return tailored


def rebuild_tailored_docx(source_resume_path: str, edited_text: str, destination_path: str) -> None:
    """Re-applies manually edited lines onto a fresh copy of the ORIGINAL resume,
    so a manual edit preserves the same fonts/colors/styling as a Claude tailoring
    pass — never builds a plain-text document from scratch.
    """
    document = Document(source_resume_path)
    paragraphs = document.paragraphs
    header_end_index = _detect_header_end(paragraphs)
    edited_lines = edited_text.splitlines()

    for index, paragraph in enumerate(paragraphs):
        if index <= header_end_index:
            continue
        if index < len(edited_lines):
            _replace_paragraph_text(paragraph, edited_lines[index])

    document.save(destination_path)


def _apply_tailored_lines(paragraphs, header_end_index: int, tailored_by_line_number: dict) -> list[str]:
    final_lines = []
    for index, paragraph in enumerate(paragraphs):
        if index <= header_end_index:
            final_lines.append(paragraph.text)
            continue

        new_text = tailored_by_line_number.get(index + 1)
        if new_text is None:
            final_lines.append(paragraph.text)
            continue

        _replace_paragraph_text(paragraph, new_text)
        final_lines.append(new_text)

    return final_lines


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replaces only the text content of a paragraph's runs, never the paragraph
    itself — this is what keeps the original font/size/bold/color untouched.
    Multi-run paragraphs put all new text in the first run (its formatting wins)
    and blank out the rest, since text boundaries don't survive rewording.
    """
    runs = paragraph.runs
    if not runs:
        if new_text:
            paragraph.add_run(new_text)
        return

    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def _detect_header_end(paragraphs) -> int:
    """Returns the 0-based index of the last header line (name/headline/contact).

    Heuristic: scan the first few lines for an email/phone/LinkedIn/GitHub
    pattern — the header ends at the last such line found. If none is found,
    fall back to treating the first two lines (name + headline) as the header,
    since those must never be touched even when contact info can't be detected.
    """
    scan_limit = min(HEADER_SCAN_LINES, len(paragraphs))
    last_contact_index = None
    for index in range(scan_limit):
        if CONTACT_LINE_PATTERN.search(paragraphs[index].text):
            last_contact_index = index

    if last_contact_index is not None:
        return last_contact_index
    return min(1, len(paragraphs) - 1)


def _numbered_text_within_budget(lines: list[str], char_budget: int) -> tuple[str, int]:
    """Builds a numbered-line block, including only whole lines that fit the
    budget — never truncates mid-line, since that would break the line-number
    contract the round-trip with Claude depends on. Lines past the cutoff are
    simply never sent, and stay unchanged in the final document.
    """
    included = []
    total = 0
    for index, line in enumerate(lines):
        numbered_line = f"{index + 1}. {line}"
        addition = len(numbered_line) + 1
        if included and total + addition > char_budget:
            break
        included.append(numbered_line)
        total += addition
    return "\n".join(included), len(included)


def _build_file_name(company: str, job_title: str) -> str:
    title_part = _format_job_title(job_title)
    company_part = _format_company(company)
    date_part = date.today().isoformat()
    return f"{title_part}_{company_part}_{date_part}.docx"


def _format_job_title(job_title: str) -> str:
    words = re.findall(r"[A-Za-z0-9]+", job_title or "")
    kept = [word for word in words if word.lower() not in JOB_TITLE_REMOVE_WORDS]
    pascal_case = "".join(word.capitalize() for word in kept)
    return (pascal_case or "Role")[:JOB_TITLE_MAX_LENGTH]


def _format_company(company: str) -> str:
    words = re.findall(r"[A-Za-z0-9]+", company or "")
    kept = [word for word in words if word.lower() not in COMPANY_REMOVE_WORDS]
    joined = "".join(kept)
    return (joined or "Company")[:COMPANY_MAX_LENGTH]
