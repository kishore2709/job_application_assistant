from pathlib import Path

from docx import Document

from app.db.repositories import ResumeRepository
from app.models.resume import Resume
from app.utils.file_utils import copy_resume_to_default


class InvalidResumeFileError(Exception):
    pass


class ResumeCopyError(Exception):
    pass


class ResumeService:
    def __init__(self) -> None:
        self.repository = ResumeRepository()

    def upload_resume(self, file_path: str) -> Resume:
        source = Path(file_path)
        if source.suffix.lower() != ".docx":
            raise InvalidResumeFileError("Please select a .docx file.")

        try:
            destination = copy_resume_to_default(file_path)
        except OSError as error:
            raise ResumeCopyError(f"Could not copy the resume file: {error}") from error

        resume = Resume(
            file_name=destination.name,
            file_path=str(destination),
            is_default=True,
        )
        resume_id = self.repository.add(resume)
        return self.repository.get_by_id(resume_id)

    def get_default_resume(self) -> Resume | None:
        return self.repository.get_default()

    def get_resume_text(self, file_path: str) -> str:
        try:
            document = Document(file_path)
        except Exception as error:
            raise InvalidResumeFileError(f"Could not read this resume: {error}") from error

        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)

    def set_default_resume(self, resume_id: int) -> None:
        self.repository.set_default(resume_id)

    def delete_resume(self, resume_id: int) -> None:
        resume = self.repository.get_by_id(resume_id)
        if resume is None:
            return
        Path(resume.file_path).unlink(missing_ok=True)
        self.repository.delete(resume_id)

    def list_resumes(self) -> list[Resume]:
        return self.repository.list_all()
